from __future__ import annotations

import html as html_module
import re
from pathlib import Path

from testpilot.domain.api import ApiDocument, ApiEndpoint, ApiParameter

HTTP_LINE = re.compile(r"\b(GET|POST|PUT|PATCH|DELETE|HEAD|OPTIONS)\s+((?:https?://[^\s]+)|(?:/[^\s`|，。；;]*))", re.I)


class DocumentParser:
    """Extract conservative endpoint drafts; every result requires manual review."""

    def parse_file(self, path: str | Path) -> ApiDocument:
        source = Path(path)
        suffix = source.suffix.lower()
        if suffix in {".md", ".txt"}:
            text = source.read_text(encoding="utf-8-sig")
        elif suffix in {".html", ".htm"}:
            raw = source.read_text(encoding="utf-8-sig")
            text = html_module.unescape(re.sub(r"<[^>]+>", " ", raw))
        elif suffix == ".docx":
            from docx import Document
            document = Document(source)
            text = "\n".join(p.text for p in document.paragraphs)
            text += "\n" + "\n".join(" | ".join(cell.text for cell in row.cells) for t in document.tables for row in t.rows)
        elif suffix == ".pdf":
            from pypdf import PdfReader
            text = "\n".join(page.extract_text() or "" for page in PdfReader(source).pages)
        elif suffix in {".xlsx", ".xlsm"}:
            return self._parse_excel(source)
        else:
            raise ValueError(f"不支持的文档格式：{suffix}")
        return self.parse_text(text, source.name)

    def parse_text(self, text: str, source_name: str = "粘贴文档") -> ApiDocument:
        endpoints, bases, seen = [], set(), set()
        for line_number, line in enumerate(text.splitlines(), 1):
            for match in HTTP_LINE.finditer(line):
                method, raw_url = match.group(1).upper(), match.group(2).rstrip(".,;，。；")
                base_url, path = _split_url(raw_url)
                if base_url:
                    bases.add(base_url)
                if (method, path) in seen:
                    continue
                seen.add((method, path))
                summary = line[:match.start()].strip(" #|:-") or "文档提取接口"
                endpoints.append(ApiEndpoint(
                    method, path, summary[:200], module="文档草稿", source="document",
                    source_location=f"{source_name}:{line_number}",
                ))
        return ApiDocument(Path(source_name).stem, "", "Document draft", sorted(bases), endpoints)

    def _parse_excel(self, source: Path) -> ApiDocument:
        from openpyxl import load_workbook

        workbook = load_workbook(source, read_only=True, data_only=True)
        endpoints = []
        aliases = {
            "method": {"method", "方法", "请求方法"},
            "path": {"path", "url", "路径", "接口地址", "请求地址"},
            "summary": {"summary", "name", "名称", "接口名称", "描述"},
            "module": {"module", "模块", "目录"},
        }
        for sheet in workbook.worksheets:
            rows = sheet.iter_rows(values_only=True)
            headers = next(rows, None)
            if not headers:
                continue
            normalized = {str(value).strip().lower(): index for index, value in enumerate(headers) if value is not None}
            mapping = {
                key: next((normalized[name] for name in names if name in normalized), None) for key, names in aliases.items()
            }
            if mapping["method"] is None or mapping["path"] is None:
                continue
            for row_number, row in enumerate(rows, 2):
                method = str(row[mapping["method"]] or "").upper()
                raw_url = str(row[mapping["path"]] or "")
                if method not in {"GET", "POST", "PUT", "PATCH", "DELETE", "HEAD", "OPTIONS"} or not raw_url:
                    continue
                _, path = _split_url(raw_url)
                endpoints.append(ApiEndpoint(
                    method, path,
                    str(row[mapping["summary"]] or "") if mapping["summary"] is not None else "",
                    module=str(row[mapping["module"]] or sheet.title) if mapping["module"] is not None else sheet.title,
                    source="document", source_location=f"{source.name}:{sheet.title}!{row_number}",
                ))
        return ApiDocument(source.stem, "", "Excel draft", [], endpoints)


def _split_url(value: str) -> tuple[str, str]:
    from urllib.parse import urlsplit

    if value.startswith(("http://", "https://")):
        parsed = urlsplit(value)
        return f"{parsed.scheme}://{parsed.netloc}", parsed.path or "/"
    return "", value.split("?")[0] or "/"

